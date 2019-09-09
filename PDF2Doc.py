import PyPDF2
import docx
pdfFileObj = open('c:/Users/user_name/Location of PDF/PDF2copy.pdf', 'rb')
pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
pdfReader.numPages
doc = docx.Document()
doc.save('c:/Users/user_name/Location of new Doc/New.docx')
for pageNum in range(0, pdfReader.numPages):
    pageObj = pdfReader.getPage(pageNum)
    pageObj.extractText()
    print(pageObj.extractText())
    doc = docx.Document('c:/Users/user_name/Location of new Doc/New.docx')
    doc.add_paragraph(pageObj.extractText())
    doc.save('c:/Users/user_name/Location of new Doc/New.docx')
doc.save('c:/Users/user_name/Location of new Doc/New.docx')
input("Press Enter to exit")
quit()
